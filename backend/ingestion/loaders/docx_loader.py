"""
DOCX loader — brings .docx ingestion to parity with the PDF pipeline.

Three capabilities mirror the PDF loader so a Word document is handled as richly:
  • TABLES   → rendered as Markdown ([TABLE n] ... [/TABLE n]) so retrieval and the
               frontend treat them like PDF tables (header row + separator).
  • MATH     → Word stores equations as OMML (Office Math Markup), which python-docx's
               ``.text`` silently drops. We convert OMML to LaTeX inline (\\( ... \\) /
               \\[ ... \\]) so formulas survive — this is EXACT structured extraction, no
               vision/API needed, so it always runs.
  • FIGURES  → embedded images are captioned by the vision model (the SAME VisionCaptioner
               the PDF loader uses) and emitted as their own searchable chunks. Gated by
               VISION_ENABLED + VISION_CAPTION_FIGURES; the image calls run in parallel.
"""

from typing import List, Optional, Tuple
from langchain_core.documents import Document

from backend.config import VISION_ENABLED, VISION_CAPTION_FIGURES, VISION_MAX_WORKERS
from backend.processing.parallel_executor import ParallelExecutor
from backend.utils.logger import get_logger

log = get_logger("docx_loader")

# Office Open XML namespaces (full URIs — matched directly so we never depend on a prefix).
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"      # word text
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"        # math (OMML)
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"            # drawing
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# n-ary OMML operators → LaTeX.
_NARY_OPS = {
    "∑": r"\sum", "∏": r"\prod", "∐": r"\coprod",
    "∫": r"\int", "∬": r"\iint", "∭": r"\iiint", "∮": r"\oint",
    "⋃": r"\bigcup", "⋂": r"\bigcap", "⋁": r"\bigvee", "⋀": r"\bigwedge",
}

# Skip embedded images smaller than this on EITHER pixel dimension — bullet icons, rules,
# logos. Pixel size (not byte size) is the right gate: a clean line diagram (graph, chart)
# compresses to a TINY PNG yet is a real figure, while an icon is small in pixels. Reused
# for any format whose dimensions we can read from the header.
_DOCX_MIN_IMAGE_PX = 80
# Fallback byte floor ONLY when dimensions can't be parsed (unknown/odd encoding) — keep
# anything non-trivial rather than silently dropping a readable image.
_DOCX_MIN_IMAGE_BYTES = 1024
# Vision can read these; EMF/WMF (Word's vector fallbacks) are skipped.
_VISION_IMAGE_TYPES = ("image/png", "image/jpeg", "image/jpg", "image/gif", "image/webp")


class DOCXLoader:

    # ──────────────────────────────────────────────────────────────────
    #  XML helpers
    # ──────────────────────────────────────────────────────────────────
    @staticmethod
    def _ln(el) -> str:
        """Local (namespace-stripped) tag name of an lxml element."""
        tag = el.tag
        return tag.rsplit("}", 1)[-1] if isinstance(tag, str) else ""

    @staticmethod
    def _attr(el, name: str) -> Optional[str]:
        """Read an attribute by local name, ignoring its namespace prefix."""
        for k, v in el.attrib.items():
            if k.rsplit("}", 1)[-1] == name:
                return v
        return None

    @staticmethod
    def _clean(s: str) -> str:
        """Collapse whitespace/newlines in a cell or run to a single clean line."""
        return " ".join((s or "").replace("\n", " ").split()).strip()

    # ──────────────────────────────────────────────────────────────────
    #  OMML math → LaTeX  (structured, no vision/API)
    # ──────────────────────────────────────────────────────────────────
    @staticmethod
    def _omml_children(el) -> str:
        """Concatenate LaTeX of all child elements, skipping property nodes (*Pr)."""
        out = []
        for c in el:
            if DOCXLoader._ln(c).endswith("Pr"):  # rPr, fPr, naryPr, ctrlPr, ...
                continue
            out.append(DOCXLoader._omml_to_latex(c))
        return "".join(out)

    @staticmethod
    def _omml_part(el, name: str) -> str:
        """LaTeX of the first child with the given local name (e, num, den, sup, ...)."""
        for c in el:
            if DOCXLoader._ln(c) == name:
                return DOCXLoader._omml_children(c)
        return ""

    @staticmethod
    def _nary_op(el) -> str:
        """The operator of an n-ary (m:nary): sum/integral/product/... default integral."""
        for pr in el:
            if DOCXLoader._ln(pr) == "naryPr":
                for c in pr:
                    if DOCXLoader._ln(c) == "chr":
                        val = DOCXLoader._attr(c, "val")
                        return _NARY_OPS.get(val, val or r"\int")
        return r"\int"  # OMML default n-ary char is the integral

    @staticmethod
    def _delim_chars(el) -> Tuple[str, str]:
        """Begin/end characters of a delimiter (m:d); default round brackets."""
        beg, end = "(", ")"
        for pr in el:
            if DOCXLoader._ln(pr) == "dPr":
                for c in pr:
                    if DOCXLoader._ln(c) == "begChr":
                        beg = DOCXLoader._attr(c, "val") or beg
                    elif DOCXLoader._ln(c) == "endChr":
                        end = DOCXLoader._attr(c, "val") or end
        return beg, end

    @staticmethod
    def _omml_to_latex(el) -> str:
        """Recursively convert an OMML element subtree to a LaTeX string (best-effort)."""
        ln = DOCXLoader._ln(el)

        if ln == "t":  # math text run content
            return el.text or ""
        if ln == "f":  # fraction
            return "\\frac{" + DOCXLoader._omml_part(el, "num") + "}{" + DOCXLoader._omml_part(el, "den") + "}"
        if ln == "sSup":  # superscript
            return "{" + DOCXLoader._omml_part(el, "e") + "}^{" + DOCXLoader._omml_part(el, "sup") + "}"
        if ln == "sSub":  # subscript
            return "{" + DOCXLoader._omml_part(el, "e") + "}_{" + DOCXLoader._omml_part(el, "sub") + "}"
        if ln == "sSubSup":  # sub + super
            return ("{" + DOCXLoader._omml_part(el, "e") + "}_{" + DOCXLoader._omml_part(el, "sub")
                    + "}^{" + DOCXLoader._omml_part(el, "sup") + "}")
        if ln == "rad":  # radical
            deg = DOCXLoader._omml_part(el, "deg")
            e = DOCXLoader._omml_part(el, "e")
            return ("\\sqrt[" + deg + "]{" + e + "}") if deg else ("\\sqrt{" + e + "}")
        if ln == "nary":  # sum / integral / product
            op = DOCXLoader._nary_op(el)
            sub = DOCXLoader._omml_part(el, "sub")
            sup = DOCXLoader._omml_part(el, "sup")
            e = DOCXLoader._omml_part(el, "e")
            s = op
            if sub:
                s += "_{" + sub + "}"
            if sup:
                s += "^{" + sup + "}"
            return s + " " + e
        if ln == "d":  # delimiter (...) [...] {...}
            beg, end = DOCXLoader._delim_chars(el)
            return "\\left" + beg + " " + DOCXLoader._omml_children(el) + " \\right" + end
        if ln == "func":  # named function: sin, lim, ...
            return DOCXLoader._omml_part(el, "fName") + "(" + DOCXLoader._omml_part(el, "e") + ")"
        if ln == "limLow":
            return "{" + DOCXLoader._omml_part(el, "e") + "}_{" + DOCXLoader._omml_part(el, "lim") + "}"
        if ln == "limUpp":
            return "{" + DOCXLoader._omml_part(el, "e") + "}^{" + DOCXLoader._omml_part(el, "lim") + "}"
        if ln == "bar":
            return "\\overline{" + DOCXLoader._omml_part(el, "e") + "}"
        if ln == "acc":
            return "\\hat{" + DOCXLoader._omml_part(el, "e") + "}"
        if ln == "m":  # matrix
            rows = []
            for mr in el:
                if DOCXLoader._ln(mr) == "mr":
                    cells = [DOCXLoader._omml_to_latex(e) for e in mr if DOCXLoader._ln(e) == "e"]
                    rows.append(" & ".join(cells))
            return "\\begin{matrix} " + " \\\\ ".join(rows) + " \\end{matrix}"

        # Containers (oMath, e, num, den, ...) and any unknown element → process children.
        return DOCXLoader._omml_children(el)

    # ──────────────────────────────────────────────────────────────────
    #  Paragraph text (interleaving normal runs and inline/display math)
    # ──────────────────────────────────────────────────────────────────
    @staticmethod
    def _run_text(r_el) -> str:
        """Text of a word run (w:r): w:t content, tabs and breaks. Ignores math runs."""
        out = []
        for t in r_el.iter():
            if t.tag == f"{{{_W_NS}}}t":
                out.append(t.text or "")
            elif t.tag == f"{{{_W_NS}}}tab":
                out.append("\t")
            elif t.tag in (f"{{{_W_NS}}}br", f"{{{_W_NS}}}cr"):
                out.append("\n")
        return "".join(out)

    @staticmethod
    def _paragraph_text(p_el) -> str:
        """Extract a paragraph's text in document order, converting OMML math to LaTeX.

        Inline equations (m:oMath) become \\( ... \\); displayed ones (m:oMathPara) become
        \\[ ... \\]. Normal runs (incl. those inside hyperlinks) keep their text. This is
        what python-docx's plain ``.text`` cannot do — it drops the math entirely.
        """
        parts = []
        for child in p_el:
            ln = DOCXLoader._ln(child)
            if ln == "r":
                parts.append(DOCXLoader._run_text(child))
            elif ln == "hyperlink":
                for r in child:
                    if DOCXLoader._ln(r) == "r":
                        parts.append(DOCXLoader._run_text(r))
            elif ln == "oMath":
                latex = DOCXLoader._omml_to_latex(child).strip()
                if latex:
                    parts.append(" \\(" + latex + "\\) ")
            elif ln == "oMathPara":
                for m in child:
                    if DOCXLoader._ln(m) == "oMath":
                        latex = DOCXLoader._omml_to_latex(m).strip()
                        if latex:
                            parts.append(" \\[" + latex + "\\] ")
        return "".join(parts).strip()

    # ──────────────────────────────────────────────────────────────────
    #  Tables → Markdown
    # ──────────────────────────────────────────────────────────────────
    @staticmethod
    def _table_to_markdown(table) -> str:
        """Render a DOCX table as Markdown (header row + separator), like the PDF loader."""
        rows = []
        for row in table.rows:
            rows.append([DOCXLoader._clean(cell.text) for cell in row.cells])
        rows = [r for r in rows if any(c for c in r)]
        if not rows:
            return ""

        width = max(len(r) for r in rows)
        rows = [r + [""] * (width - len(r)) for r in rows]
        header = [h or f"Column_{i + 1}" for i, h in enumerate(rows[0])]

        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * width) + " |",
        ]
        for r in rows[1:]:
            lines.append("| " + " | ".join(r) + " |")
        return "\n".join(lines).strip()

    @staticmethod
    def _iter_blocks(doc):
        """Yield paragraph (text+math) and table (Markdown) blocks in document order."""
        from docx.table import Table

        table_count = 0
        for child in doc.element.body.iterchildren():
            ln = DOCXLoader._ln(child)
            if ln == "p":
                text = DOCXLoader._paragraph_text(child)
                if text:
                    yield text
            elif ln == "tbl":
                md = DOCXLoader._table_to_markdown(Table(child, doc))
                if md:
                    table_count += 1
                    yield f"[TABLE {table_count}]\n{md}\n[/TABLE {table_count}]"

    # ──────────────────────────────────────────────────────────────────
    #  Embedded images → vision captions (parallel)
    # ──────────────────────────────────────────────────────────────────
    @staticmethod
    def _image_dimensions(blob: bytes) -> Tuple[int, int]:
        """Read (width, height) in pixels from an image's header — dependency-free.

        Supports PNG, JPEG, GIF and (RIFF) WebP, which covers every type the vision model
        accepts. Returns (0, 0) when the format/dimensions can't be determined, so the
        caller can fall back to a byte-size check. No full decode, so it's cheap.
        """
        try:
            if blob[:8] == b"\x89PNG\r\n\x1a\n":
                # IHDR width/height are the two big-endian uint32 at offsets 16 and 20.
                return int.from_bytes(blob[16:20], "big"), int.from_bytes(blob[20:24], "big")
            if blob[:3] == b"\xff\xd8\xff":  # JPEG: scan for an SOF marker
                i, n = 2, len(blob)
                while i + 9 < n:
                    if blob[i] != 0xFF:
                        i += 1
                        continue
                    marker = blob[i + 1]
                    # SOF0..SOF15 carry the frame dimensions (skip non-SOF / RSTn / padding).
                    if 0xC0 <= marker <= 0xCF and marker not in (0xC4, 0xC8, 0xCC):
                        h = int.from_bytes(blob[i + 5:i + 7], "big")
                        w = int.from_bytes(blob[i + 7:i + 9], "big")
                        return w, h
                    seg_len = int.from_bytes(blob[i + 2:i + 4], "big")
                    if seg_len <= 0:
                        break
                    i += 2 + seg_len
            if blob[:6] in (b"GIF87a", b"GIF89a"):
                return int.from_bytes(blob[6:8], "little"), int.from_bytes(blob[8:10], "little")
            if blob[:4] == b"RIFF" and blob[8:12] == b"WEBP":
                fmt = blob[12:16]
                if fmt == b"VP8 ":
                    return (int.from_bytes(blob[26:28], "little") & 0x3FFF,
                            int.from_bytes(blob[28:30], "little") & 0x3FFF)
                if fmt == b"VP8L":
                    b = blob[21:25]
                    w = ((b[0] | (b[1] << 8)) & 0x3FFF) + 1
                    h = (((b[1] >> 6) | (b[2] << 2) | ((b[3] & 0x0F) << 10)) & 0x3FFF) + 1
                    return w, h
                if fmt == b"VP8X":
                    return ((int.from_bytes(blob[24:27], "little") & 0xFFFFFF) + 1,
                            (int.from_bytes(blob[27:30], "little") & 0xFFFFFF) + 1)
        except Exception:
            pass
        return 0, 0

    @staticmethod
    def _extract_images(doc) -> List[Tuple[bytes, str]]:
        """Collect embedded raster images as (blob, mime), de-duplicated and size-gated.

        Reads the drawing blips in body order and resolves each relationship to its image
        part. Vector formats (EMF/WMF) the vision model can't read are skipped. A figure is
        kept when its PIXEL dimensions clear the minimum (clean line diagrams compress to
        tiny PNGs, so byte size would wrongly drop them); when dimensions are unreadable, a
        small byte floor is used instead.
        """
        out: List[Tuple[bytes, str]] = []
        seen = set()
        for blip in doc.element.body.iter(f"{{{_A_NS}}}blip"):
            rid = blip.get(f"{{{_R_NS}}}embed") or blip.get(f"{{{_R_NS}}}link")
            if not rid or rid in seen:
                continue
            seen.add(rid)
            try:
                part = doc.part.related_parts[rid]
            except (KeyError, AttributeError):
                continue
            ctype = (getattr(part, "content_type", "") or "").lower()
            if ctype not in _VISION_IMAGE_TYPES:
                continue
            blob = getattr(part, "blob", b"")
            if not blob:
                continue
            w, h = DOCXLoader._image_dimensions(blob)
            if w and h:
                if w >= _DOCX_MIN_IMAGE_PX and h >= _DOCX_MIN_IMAGE_PX:
                    out.append((blob, ctype))
            elif len(blob) >= _DOCX_MIN_IMAGE_BYTES:
                out.append((blob, ctype))
        return out

    @staticmethod
    def _caption_task(task: dict, index: int, total: int) -> dict:
        """Parallel worker: caption one embedded image (thread-safe — only bytes)."""
        from backend.ingestion.vision_captioner import get_vision_captioner
        caption = get_vision_captioner().caption_figure(task["blob"], task["mime"])
        return {"idx": task["idx"], "caption": caption}

    @staticmethod
    def _caption_images(doc, doc_id: str, filename: str, clean_fn) -> List[Document]:
        """Detect, caption (in parallel) and wrap embedded images as their own Documents.

        Each caption becomes a searchable text chunk, mirroring the PDF figure-caption
        chunks. Returns [] when figures are off or none are found/captioned.
        """
        images = DOCXLoader._extract_images(doc)
        if not images:
            return []

        tasks = [{"idx": i, "blob": b, "mime": m} for i, (b, m) in enumerate(images)]
        log.info(f"[LOAD DOCX] {filename}: captioning {len(tasks)} image(s) in parallel "
                 f"(max_workers={VISION_MAX_WORKERS})")
        results = ParallelExecutor.execute_parallel(
            tasks=tasks,
            task_func=DOCXLoader._caption_task,
            max_workers=VISION_MAX_WORKERS,
            operation_name="docx-image-caption",
        )

        by_idx = {}
        for r in results:
            if r and not r.get("error") and (r.get("caption") or "").strip():
                by_idx[r["idx"]] = r["caption"].strip()

        figure_docs: List[Document] = []
        fig_n = 0
        for i in range(len(images)):
            caption = by_idx.get(i)
            if not caption:
                continue
            fig_n += 1
            body = f"Figure {fig_n}: {caption}"
            if clean_fn:
                body = clean_fn(body)
            if body.strip():
                figure_docs.append(Document(
                    page_content=body.strip(),
                    metadata={
                        "doc_id": doc_id, "filename": filename, "source": filename,
                        "file_type": "docx", "page": 1,
                    },
                ))
        if figure_docs:
            log.info(f"[LOAD DOCX] {filename}: captioned {len(figure_docs)} image(s) as text chunk(s)")
        return figure_docs

    # ──────────────────────────────────────────────────────────────────
    #  Load
    # ──────────────────────────────────────────────────────────────────
    @staticmethod
    def load(file_path: str, filename: str, doc_id: str, clean_fn=None) -> List[Document]:
        log.info(f"[LOAD DOCX] {filename} — paragraphs + math + tables")

        documents: List[Document] = []
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)

            # ── PASS 1: linear text (with inline LaTeX math) + Markdown tables, buffered
            #    into ~logical pages so chunk metadata stays page-like. ──
            pages: List[Document] = []
            page_num = 1
            logical_page_size = 500
            buffer: List[str] = []
            buffer_len = 0

            for block_text in DOCXLoader._iter_blocks(doc):
                buffer.append(block_text)
                buffer_len += len(block_text)

                if buffer_len >= logical_page_size:
                    content = " ".join(buffer)
                    if clean_fn:
                        content = clean_fn(content)
                    if content.strip():
                        pages.append(Document(
                            page_content=content.strip(),
                            metadata={
                                "doc_id": doc_id, "filename": filename, "source": filename,
                                "file_type": "docx", "page": int(page_num),
                            },
                        ))
                    page_num += 1
                    buffer = []
                    buffer_len = 0

            if buffer:
                content = " ".join(buffer)
                if clean_fn:
                    content = clean_fn(content)
                if content.strip():
                    pages.append(Document(
                        page_content=content.strip(),
                        metadata={
                            "doc_id": doc_id, "filename": filename, "source": filename,
                            "file_type": "docx", "page": int(page_num),
                        },
                    ))

            documents.extend(pages)

            # ── PASS 2: embedded images → vision caption chunks (parallel). ──
            if VISION_ENABLED and VISION_CAPTION_FIGURES:
                try:
                    documents.extend(DOCXLoader._caption_images(doc, doc_id, filename, clean_fn))
                except Exception as fig_err:
                    log.warning(f"[LOAD DOCX] {filename}: image captioning skipped "
                                f"({type(fig_err).__name__}: {fig_err})")

            log.info(f"[LOAD DOCX] ✅ {filename} — {len(documents)} document(s) "
                     f"({len(pages)} text page(s))")
            return [d for d in documents if d.page_content.strip()]

        except Exception as e:
            log.error(f"[LOAD DOCX] ❌ FAILED for {filename} — {type(e).__name__}: {e}", exc_info=True)
            return []
