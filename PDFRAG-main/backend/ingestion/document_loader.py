import os  
import uuid  
from typing import List, Dict, Any

from langchain_community.document_loaders import PyMuPDFLoader  
from langchain_core.documents import Document  
from backend.utils.logger import get_logger

log = get_logger("document_loader")


class DocumentLoader:  
    def __init__(self, upload_dir: str = "uploads"):  
        self.upload_dir = upload_dir  
        os.makedirs(upload_dir, exist_ok=True)  
        log.info(f"DocumentLoader initialized — upload_dir: {upload_dir}")

    def save_file(self, filename: str, content: bytes) -> str:  
        safe_name = f"{uuid.uuid4().hex}_{filename}"  
        file_path = os.path.join(self.upload_dir, safe_name)  
        with open(file_path, "wb") as f:  
            f.write(content)  
        log.info(f"[SAVE] {filename} → {file_path} ({len(content)} bytes)")  
        return file_path

    def load_pdf(self, file_path: str, filename: str, doc_id: str) -> List[Document]:  
        """Load PDF using PyMuPDF."""  
        log.info(f"[LOAD PDF] {filename}")  
        loader = PyMuPDFLoader(file_path)  
        pages = loader.load()  
        for page in pages:  
            page.metadata["doc_id"]   = doc_id  
            page.metadata["filename"] = filename  
            page.metadata["source"]   = filename  
            page.metadata["page"]     = page.metadata.get("page", 0) + 1  
            page.page_content         = self._clean(page.page_content)  
        return [p for p in pages if p.page_content.strip()]

    def load_docx(self, file_path: str, filename: str, doc_id: str) -> List[Document]:  
        """Load DOCX using python-docx."""  
        log.info(f"[LOAD DOCX] {filename}")  
        from docx import Document as DocxDocument

        doc   = DocxDocument(file_path)  
        pages = []  
        page_num = 1

        # Group paragraphs into logical "pages" of ~500 chars each  
        buffer = []  
        buffer_len = 0

        for para in doc.paragraphs:  
            text = para.text.strip()  
            if not text:  
                continue  
            buffer.append(text)  
            buffer_len += len(text)

            if buffer_len >= 500:  
                content = " ".join(buffer)  
                pages.append(Document(  
                    page_content=self._clean(content),  
                    metadata={  
                        "doc_id":   doc_id,  
                        "filename": filename,  
                        "source":   filename,  
                        "page":     page_num,  
                    }  
                ))  
                page_num  += 1  
                buffer     = []  
                buffer_len = 0

        # Remaining text  
        if buffer:  
            content = " ".join(buffer)  
            pages.append(Document(  
                page_content=self._clean(content),  
                metadata={  
                    "doc_id":   doc_id,  
                    "filename": filename,  
                    "source":   filename,  
                    "page":     page_num,  
                }  
            ))

        log.info(f"[LOAD DOCX] ✅ {filename} — {len(pages)} logical pages extracted")  
        return [p for p in pages if p.page_content.strip()]

    def load_pdf_or_docx(self, file_path: str, filename: str) -> Dict[str, Any]:  
        doc_id = uuid.uuid4().hex  
        try:  
            ext = filename.lower().split(".")[-1]

            if ext == "pdf":  
                pages = self.load_pdf(file_path, filename, doc_id)  
            elif ext in ("docx", "doc"):  
                pages = self.load_docx(file_path, filename, doc_id)  
            else:  
                raise ValueError(f"Unsupported file type: .{ext}")

            log.info(f"[LOAD] ✅ {filename} — {len(pages)} pages | {sum(len(p.page_content) for p in pages)} chars")  
            return {  
                "doc_id":      doc_id,  
                "filename":    filename,  
                "total_pages": len(pages),  
                "total_chars": sum(len(p.page_content) for p in pages),  
                "documents":   pages,  
                "errors":      [],  
            }

        except Exception as e:  
            log.error(f"[LOAD] ❌ {filename} — {e}", exc_info=True)  
            return {  
                "doc_id":      doc_id,  
                "filename":    filename,  
                "total_pages": 0,  
                "total_chars": 0,  
                "documents":   [],  
                "errors":      [str(e)],  
            }

    def load_documents(self, files: List[tuple]) -> List[Dict[str, Any]]:  
        results = []  
        for filename, content in files:  
            file_path = self.save_file(filename, content)  
            results.append(self.load_pdf_or_docx(file_path, filename))  
        return results

    @staticmethod  
    def _clean(text: str) -> str:  
        lines = [l.strip() for l in text.splitlines() if l.strip()]  
        return " ".join(lines).strip()  
