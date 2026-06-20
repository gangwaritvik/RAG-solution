from typing import List  
from langchain_core.documents import Document  
from backend.utils.logger import get_logger

log = get_logger("docx_loader")


class DOCXLoader:

    @staticmethod
    def _table_to_text(table) -> str:
        """Convert a DOCX table into plain text rows for embedding/retrieval."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        return "\n".join(rows).strip()

    @staticmethod
    def _iter_blocks(doc):
        """Yield paragraph/table blocks in original document order."""
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        body = doc.element.body
        for child in body.iterchildren():
            tag = child.tag.lower()
            if tag.endswith("}p"):
                para = Paragraph(child, doc)
                text = para.text.strip()
                if text:
                    yield text
            elif tag.endswith("}tbl"):
                table = Table(child, doc)
                table_text = DOCXLoader._table_to_text(table)
                if table_text:
                    yield f"[TABLE]\n{table_text}\n[/TABLE]"

    @staticmethod  
    def load(file_path: str, filename: str, doc_id: str, clean_fn=None) -> List[Document]:  
        log.info(f"[LOAD DOCX] {filename} — paragraphs + tables")

        try:  
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)

            pages: List[Document] = []  
            page_num = 1  
            logical_page_size = 500

            buffer = []  
            buffer_len = 0
            table_count = 0

            for block_text in DOCXLoader._iter_blocks(doc):
                text = block_text.strip()
                if text.startswith("[TABLE]"):
                    table_count += 1

                buffer.append(text)  
                buffer_len += len(text)

                if buffer_len >= logical_page_size:  
                    content = " ".join(buffer)  
                    if clean_fn:  
                        content = clean_fn(content)

                    if content:  
                        pages.append(  
                            Document(  
                                page_content=content,  
                                metadata={  
                                    "doc_id": doc_id,  
                                    "filename": filename,  
                                    "source": filename,  
                                    "file_type": "docx",  
                                    "page": int(page_num),  
                                },  
                            )  
                        )

                    page_num += 1  
                    buffer = []  
                    buffer_len = 0

            # Remaining text  
            if buffer:  
                content = " ".join(buffer)  
                if clean_fn:  
                    content = clean_fn(content)

                if content:  
                    pages.append(  
                        Document(  
                            page_content=content,  
                            metadata={  
                                "doc_id": doc_id,  
                                "filename": filename,  
                                "source": filename,  
                                "file_type": "docx",  
                                "page": int(page_num),  
                            },  
                        )  
                    )

            log.info(
                f"[LOAD DOCX] ✅ {filename} — {len(pages)} logical pages extracted "
                f"(tables={table_count})"
            )
            return [p for p in pages if p.page_content.strip()]

        except Exception as e:  
            log.error(f"[LOAD DOCX] ❌ FAILED for {filename} — {type(e).__name__}: {e}", exc_info=True)  
            return []  
